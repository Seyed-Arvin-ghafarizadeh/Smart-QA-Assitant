"""Document validation utilities."""
import os
from typing import Dict, Any, Optional
from pathlib import Path

from app.exceptions import (
    FileTypeNotSupportedError,
    FileSizeExceededError,
    DocumentCorruptedError,
    DocumentEmptyError,
    PageLimitExceededError,
    ServiceUnavailableError
)

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentValidator:
    """Base class for document validators."""

    SUPPORTED_EXTENSIONS = []

    @classmethod
    def validate_file_type(cls, filename: str) -> str:
        """Validate file type and return clean extension."""
        if not filename:
            raise FileTypeNotSupportedError("File name is required.")

        file_extension = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
        full_extension = f'.{file_extension}'

        if full_extension not in cls.SUPPORTED_EXTENSIONS:
            raise FileTypeNotSupportedError(
                f"Unsupported file type. Supported formats: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )

        return full_extension

    @classmethod
    def validate_file_size(cls, file_size_bytes: int, max_size_mb: float) -> None:
        """Validate file size."""
        file_size_mb = file_size_bytes / (1024 * 1024)
        if file_size_mb > max_size_mb:
            raise FileSizeExceededError(
                f"File size ({file_size_mb:.2f} MB) exceeds maximum allowed size ({max_size_mb} MB)."
            )


class PDFValidator(DocumentValidator):
    """Validator for PDF files."""

    SUPPORTED_EXTENSIONS = ['.pdf']

    @classmethod
    def validate_content(cls, file_path: str, settings: Dict[str, Any]) -> int:
        """
        Validate PDF content and return page count.

        Args:
            file_path: Path to PDF file
            settings: Application settings with validation limits

        Returns:
            Number of pages in the document

        Raises:
            ServiceUnavailableError: If pdfplumber is not available
            DocumentCorruptedError: If PDF is corrupted
            DocumentEmptyError: If PDF has no pages
            PageLimitExceededError: If PDF exceeds page limit
        """
        if not PDFPLUMBER_AVAILABLE:
            raise ServiceUnavailableError("pdfplumber is not available. PDF processing requires pdfplumber.")

        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)

                # Check if PDF has pages
                if total_pages == 0:
                    raise DocumentEmptyError("PDF contains no pages. Please provide a valid PDF with content.")

                # Validate page count
                if total_pages > settings.max_pages:
                    raise PageLimitExceededError(
                        f"PDF has {total_pages} pages, which exceeds the maximum of {settings.max_pages} pages."
                    )

                # Try to extract text from first page to verify readability
                if pdf.pages:
                    first_page = pdf.pages[0]
                    first_page_text = first_page.extract_text() or ""
                    if len(first_page_text.strip()) == 0:
                        # Log warning but don't fail - some PDFs might be image-based
                        pass

                return total_pages

        except (DocumentEmptyError, PageLimitExceededError):
            raise
        except Exception as e:
            raise DocumentCorruptedError(f"Invalid or corrupted PDF file: {str(e)}")


class DOCXValidator(DocumentValidator):
    """Validator for DOCX files."""

    SUPPORTED_EXTENSIONS = ['.docx']

    @classmethod
    def validate_content(cls, file_path: str, settings: Dict[str, Any]) -> int:
        """
        Validate DOCX content.

        Args:
            file_path: Path to DOCX file
            settings: Application settings (not used for DOCX validation)

        Returns:
            Always returns 1 (DOCX files don't have pages)

        Raises:
            ServiceUnavailableError: If python-docx is not available
            DocumentCorruptedError: If DOCX is corrupted
            DocumentEmptyError: If DOCX has no content
        """
        if not DOCX_AVAILABLE:
            raise ServiceUnavailableError("python-docx is not available. DOCX processing requires python-docx.")

        try:
            doc = DocxDocument(file_path)
            # Basic validation - check if document has any content
            has_content = bool(doc.paragraphs or doc.tables)
            if not has_content:
                raise DocumentEmptyError("DOCX file appears to be empty or has no readable content.")
            return 1  # DOCX doesn't have pages

        except (DocumentEmptyError,):
            raise
        except Exception as e:
            raise DocumentCorruptedError(f"Invalid or corrupted DOCX file: {str(e)}")


class TXTValidator(DocumentValidator):
    """Validator for TXT files."""

    SUPPORTED_EXTENSIONS = ['.txt']

    @classmethod
    def validate_content(cls, file_path: str, settings: Dict[str, Any]) -> int:
        """
        Validate TXT content.

        Args:
            file_path: Path to TXT file
            settings: Application settings with max_characters limit

        Returns:
            Always returns 1 (TXT files don't have pages)

        Raises:
            DocumentCorruptedError: If file cannot be read
            DocumentEmptyError: If file is empty
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

                if not content.strip():
                    raise DocumentEmptyError("TXT file is empty or contains no readable text.")

                # Check file size for text files
                max_text_chars = getattr(settings, 'max_characters', 2000000)
                if len(content) > max_text_chars:
                    raise FileSizeExceededError(
                        f"TXT file is too large ({len(content):,} characters). "
                        f"Maximum allowed: {max_text_chars:,} characters."
                    )

            return 1  # TXT doesn't have pages

        except (DocumentEmptyError, FileSizeExceededError):
            raise
        except Exception as e:
            raise DocumentCorruptedError(f"Invalid TXT file: {str(e)}")


def get_validator_for_file_type(file_extension: str) -> DocumentValidator:
    """
    Get the appropriate validator for a file extension.

    Args:
        file_extension: File extension (e.g., '.pdf', '.docx', '.txt')

    Returns:
        Validator class for the file type

    Raises:
        FileTypeNotSupportedError: If file type is not supported
    """
    validators = {
        '.pdf': PDFValidator,
        '.docx': DOCXValidator,
        '.txt': TXTValidator,
    }

    validator_class = validators.get(file_extension.lower())
    if not validator_class:
        raise FileTypeNotSupportedError(f"No validator available for file type: {file_extension}")

    return validator_class


def validate_document(file_path: str, file_size_bytes: int, settings: Dict[str, Any]) -> Dict[str, Any]:
    """
    Comprehensive document validation.

    Args:
        file_path: Path to the document file
        file_size_bytes: Size of the file in bytes
        settings: Application settings with validation limits

    Returns:
        Dict with validation results including page_count

    Raises:
        Various validation errors
    """
    filename = Path(file_path).name

    # Get validator for file type
    file_extension = get_validator_for_file_type(Path(file_path).suffix)

    # Validate file type
    file_extension.validate_file_type(filename)

    # Validate file size
    file_extension.validate_file_size(file_size_bytes, settings.max_file_size_mb)

    # Validate content
    page_count = file_extension.validate_content(file_path, settings)

    return {
        'file_extension': file_extension,
        'page_count': page_count,
        'filename': filename
    }
