"""Document processing service for PDF, DOCX, and TXT extraction and chunking."""
import os
import re
from typing import List, Tuple, Optional
from pathlib import Path

from app.models.document import Chunk
from app.utils.logger import logger
from app.utils.text_cleaner import clean_text
from app.exceptions import ExtractionError, ServiceUnavailableError
from app.services.ocr_service import EasyOCRService

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber is not available. PDF processing will not work.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx is not available. DOCX processing will not work.")


def get_file_type(file_path: str) -> str:
    """Determine file type based on file extension."""
    file_path = Path(file_path)
    extension = file_path.suffix.lower()

    if extension == '.pdf':
        return 'pdf'
    elif extension == '.docx':
        return 'docx'
    elif extension == '.txt':
        return 'txt'
    else:
        raise ValueError(f"Unsupported file type: {extension}. Supported: .pdf, .docx, .txt")


def extract_text_from_pdf(
    file_path: str,
    ocr_service: Optional[EasyOCRService] = None,
    ocr_enabled: bool = True,
    ocr_text_threshold: int = 50,
    ocr_dpi: int = 300,
) -> List[Tuple[int, str]]:
    """
    Extract text from PDF using pdfplumber with OCR fallback.

    Args:
        file_path: Path to PDF file
        ocr_service: Optional OCR service instance
        ocr_enabled: Whether OCR fallback is enabled
        ocr_text_threshold: Minimum characters before OCR fallback (default: 50)
        ocr_dpi: DPI for OCR image conversion (default: 300)

    Returns:
        List of (page_number, page_text) tuples

    Raises:
        ServiceUnavailableError: If pdfplumber is not available
        ExtractionError: If PDF processing fails
    """
    if not PDFPLUMBER_AVAILABLE:
        raise ServiceUnavailableError("pdfplumber is not available. Please install pdfplumber.")

    pages_data = []
    pages_needing_ocr = []

    # First pass: Try pdfplumber extraction
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    text = page.extract_text() or ""
                    text = clean_text(text)
                    text_length = len(text.strip())

                    # Check if text extraction is sufficient
                    if text_length < ocr_text_threshold and ocr_enabled:
                        # Mark page for OCR processing
                        pages_needing_ocr.append(page_num)
                        logger.info(
                            f"Page {page_num} has insufficient text ({text_length} chars < {ocr_text_threshold}), "
                            "will use OCR fallback"
                        )
                        pages_data.append((page_num, text))  # Keep partial text
                    else:
                        pages_data.append((page_num, text))
                except Exception as e:
                    logger.warning(f"Error extracting text from PDF page {page_num}: {str(e)}")
                    if ocr_enabled:
                        pages_needing_ocr.append(page_num)
                    pages_data.append((page_num, ""))
    except Exception as e:
        logger.error(f"Error opening PDF file {file_path}: {str(e)}")
        raise ExtractionError(f"Failed to process PDF file: {str(e)}")

    # Second pass: Use OCR for pages that need it
    if pages_needing_ocr and ocr_enabled and ocr_service:
        if not ocr_service.is_available():
            logger.warning(
                "OCR service not available. Some pages may have insufficient text extraction."
            )
        else:
            logger.info(f"Processing {len(pages_needing_ocr)} pages with OCR...")
            ocr_results = ocr_service.extract_text_from_pdf(
                file_path, page_numbers=pages_needing_ocr, dpi=ocr_dpi
            )

            # Update pages_data with OCR results
            ocr_dict = {page_num: (text, confidence) for page_num, text, confidence in ocr_results}
            updated_pages_data = []

            for page_num, existing_text in pages_data:
                if page_num in ocr_dict:
                    ocr_text, ocr_confidence = ocr_dict[page_num]
                    # Merge OCR text with existing text (if any)
                    if existing_text.strip():
                        combined_text = f"{existing_text}\n\n{ocr_text}"
                    else:
                        combined_text = ocr_text

                    combined_text = clean_text(combined_text)
                    updated_pages_data.append((page_num, combined_text))
                    logger.info(
                        f"Page {page_num} OCR completed: {len(combined_text)} chars "
                        f"(confidence: {ocr_confidence:.2%})"
                    )
                else:
                    updated_pages_data.append((page_num, existing_text))

            pages_data = updated_pages_data

    return pages_data


def extract_text_from_docx(file_path: str) -> List[Tuple[int, str]]:
    """
    Extract text from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        List of (page_number, page_text) tuples (DOCX doesn't have pages, so page_number=1)

    Raises:
        ServiceUnavailableError: If python-docx is not available
        ExtractionError: If DOCX processing fails
    """
    if not DOCX_AVAILABLE:
        raise ServiceUnavailableError("python-docx is not available. Please install python-docx.")

    try:
        doc = DocxDocument(file_path)
        full_text = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        # Join all text with double newlines to simulate page breaks
        combined_text = '\n\n'.join(full_text)
        return [(1, clean_text(combined_text))]

    except Exception as e:
        logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
        raise ExtractionError(f"Failed to process DOCX file: {str(e)}")


def extract_text_from_txt(file_path: str) -> List[Tuple[int, str]]:
    """
    Extract text from TXT file.

    Args:
        file_path: Path to TXT file

    Returns:
        List of (page_number, page_text) tuples (TXT doesn't have pages, so page_number=1)

    Raises:
        ExtractionError: If file reading fails
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return [(1, clean_text(text))]
    except Exception as e:
        logger.error(f"Error reading TXT file {file_path}: {str(e)}")
        raise ExtractionError(f"Failed to process TXT file: {str(e)}")


def extract_text_from_file(file_path: str) -> List[Tuple[int, str]]:
    """
    Extract text from any supported file type.

    Args:
        file_path: Path to file

    Returns:
        List of (page_number, page_text) tuples

    Raises:
        ValueError: If file type is not supported
        ExtractionError: If extraction fails
    """
    file_type = get_file_type(file_path)

    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    elif file_type == 'txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")




class DocumentProcessor:
    """Handles document extraction, cleaning, and chunking for PDF, DOCX, and TXT files."""

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 100,
        max_workers: int = 0,
        ocr_service: Optional[EasyOCRService] = None,
        ocr_enabled: bool = True,
        ocr_text_threshold: int = 50,
        ocr_dpi: int = 300,
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Target size for text chunks (in characters)
            chunk_overlap: Overlap between chunks (in characters)
            max_workers: Max workers for parallel processing (0 = auto based on CPU cores)
            ocr_service: Optional OCR service instance for PDF OCR fallback
            ocr_enabled: Whether OCR fallback is enabled
            ocr_text_threshold: Minimum characters before OCR fallback
            ocr_dpi: DPI for OCR image conversion
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.current_chapter = None  # Track current chapter number
        self.ocr_service = ocr_service
        self.ocr_enabled = ocr_enabled
        self.ocr_text_threshold = ocr_text_threshold
        self.ocr_dpi = ocr_dpi

        # Configure max workers (default: use all CPU cores, up to 16 for better performance)
        available_cores = os.cpu_count() or 1
        self.max_workers = max_workers if max_workers > 0 else min(available_cores * 2, 16)
        logger.info(
            f"DocumentProcessor initialized with {self.max_workers} workers "
            f"(available cores: {available_cores}, OCR enabled: {ocr_enabled})"
        )

    def detect_chapter_number(self, text: str) -> Optional[int]:
        """
        Detect chapter number from text by looking for common chapter patterns.

        Args:
            text: Page text to analyze

        Returns:
            Chapter number if detected, None otherwise
        """
        if not text:
            return None

        # Common chapter patterns (case-insensitive)
        patterns = [
            r'chapter\s+(\d+)',  # "Chapter 1", "Chapter 10"
            r'chapter\s+([IVX]+)',  # "Chapter I", "Chapter IV"
            r'ch\.\s*(\d+)',  # "Ch. 1", "Ch. 10"
            r'chapter\s+(\d+)\s*:',  # "Chapter 1:"
            r'^(\d+)\s*\.\s*chapter',  # "1. Chapter"
            r'part\s+(\d+)',  # "Part 1" (sometimes used as chapters)
        ]

        # Check first 500 characters (usually where chapter headers appear)
        text_sample = text[:500].lower()

        for pattern in patterns:
            match = re.search(pattern, text_sample, re.IGNORECASE)
            if match:
                chapter_str = match.group(1)
                # Try to convert Roman numerals to numbers
                if chapter_str.isdigit():
                    return int(chapter_str)
                # Handle Roman numerals (basic)
                roman_map = {'i': 1, 'ii': 2, 'iii': 3, 'iv': 4, 'v': 5,
                           'vi': 6, 'vii': 7, 'viii': 8, 'ix': 9, 'x': 10}
                if chapter_str.lower() in roman_map:
                    return roman_map[chapter_str.lower()]

        return None

    def extract_text_from_file(self, file_path: str) -> List[Tuple[int, str, Optional[int]]]:
        """
        Extract text from any supported file type with OCR fallback for PDFs.

        Args:
            file_path: Path to file

        Returns:
            List of (page_number, page_text, chapter_number) tuples
        """
        file_type = get_file_type(file_path)

        # Use OCR-aware extraction for PDFs
        if file_type == 'pdf':
            pages_data = extract_text_from_pdf(
                file_path,
                ocr_service=self.ocr_service,
                ocr_enabled=self.ocr_enabled,
                ocr_text_threshold=self.ocr_text_threshold,
                ocr_dpi=self.ocr_dpi,
            )
        elif file_type == 'docx':
            pages_data = extract_text_from_docx(file_path)
        elif file_type == 'txt':
            pages_data = extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Convert to format with chapter detection
        result = []
        self.current_chapter = None
        total_chars_extracted = 0

        for page_num, page_text in pages_data:
            text_length = len(page_text.strip())
            total_chars_extracted += text_length

            if text_length > 0:
                # Detect chapter number from page text
                detected_chapter = self.detect_chapter_number(page_text)
                if detected_chapter is not None:
                    self.current_chapter = detected_chapter
                    logger.debug(f"Detected Chapter {detected_chapter} on page {page_num}")
            else:
                detected_chapter = None

            result.append((page_num, page_text, self.current_chapter))

        file_name = Path(file_path).name
        logger.info(
            f"Extracted text from {file_name}: {len(result)} pages/sections, "
            f"{total_chars_extracted:,} total characters extracted"
        )

        return result

    def chunk_text_simple(self, text: str) -> List[str]:
        """
        Split text into chunks using simple character-based chunking.

        Args:
            text: Text to chunk

        Returns:
            List of text chunks
        """
        if not text:
            return []

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = min(start + self.chunk_size, text_length)
            chunk = text[start:end]
            chunks.append(chunk)

            # If we've reached the end, break
            if end >= text_length:
                break

            # Move start position forward by (chunk_size - overlap) to create overlap
            start = end - self.chunk_overlap
            # Ensure we always make progress (handle edge case where overlap >= remaining text)
            if start >= end:
                start = end

        return chunks

    def chunk_text_streaming(
        self, pages_data: List[Tuple[int, str, Optional[int]]], document_id: str, filename: str
    ) -> List[Chunk]:
        """
        Split text into chunks with metadata from streaming page data.

        Args:
            pages_data: List of (page_number, page_text, chapter_number) tuples
            document_id: Unique document identifier
            filename: Original filename

        Returns:
            List of Chunk objects with metadata
        """
        chunks = []
        chunk_index = 0

        for page_num, page_text, chapter_num in pages_data:
            if not page_text.strip():
                continue

            # Chunk the page text
            page_chunks = self.chunk_text_simple(page_text)

            for chunk_text in page_chunks:
                if chunk_text.strip():  # Only create chunks with actual content
                    chunk = Chunk(
                        text=chunk_text,
                        page_number=page_num,
                        chunk_index=chunk_index,
                        document_id=document_id,
                        chapter_number=chapter_num,
                    )
                    chunks.append(chunk)
                    chunk_index += 1

        logger.info(f"Created {len(chunks)} chunks from document {document_id}")
        return chunks

    def process_document(self, file_path: str, document_id: str) -> List[Chunk]:
        """
        Process a document and return chunks with metadata.

        Args:
            file_path: Path to the document file
            document_id: Unique document identifier

        Returns:
            List of Chunk objects with metadata
        """
        filename = Path(file_path).name

        # Extract text from file
        pages_data = self.extract_text_from_file(file_path)

        # Convert to chunks
        chunks = self.chunk_text_streaming(pages_data, document_id, filename)

        return chunks
